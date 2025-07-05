"""Utility functions for the Personal Finance Assistant app."""

import os
import hashlib
import streamlit as st
from typing import List, Optional
from io import BytesIO
from pathlib import Path
import pypdf
from docx import Document


def create_directories() -> None:
    """Create necessary directories if they don't exist."""
    directories = [
        "data",
        "data/knowledge_base",
        "data/chroma_db"
    ]
    
    for directory in directories:
        os.makedirs(directory, exist_ok=True)


def get_file_hash(content: bytes) -> str:
    """Generate a hash for file content to detect duplicates."""
    return hashlib.md5(content).hexdigest()


def extract_text_from_file(uploaded_file) -> Optional[str]:
    """Extract text content from uploaded files."""
    try:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()
        
        if file_extension == '.txt':
            return str(uploaded_file.read(), "utf-8")
        
        elif file_extension == '.pdf':
            pdf_reader = pypdf.PdfReader(BytesIO(uploaded_file.read()))
            text = ""
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            return text
        
        elif file_extension == '.docx':
            doc = Document(BytesIO(uploaded_file.read()))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            st.error(f"Unsupported file type: {file_extension}")
            return None
            
    except Exception as e:
        st.error(f"Error extracting text from {uploaded_file.name}: {str(e)}")
        return None


def process_document(file_path: str) -> Optional[str]:
    """Extract text content from a file on the filesystem."""
    try:
        file_path = Path(file_path)
        
        if not file_path.exists():
            print(f"File not found: {file_path}")
            return None
        
        file_extension = file_path.suffix.lower()
        
        if file_extension == '.txt':
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        
        elif file_extension == '.pdf':
            with open(file_path, 'rb') as f:
                pdf_reader = pypdf.PdfReader(f)
                text = ""
                for page in pdf_reader.pages:
                    text += page.extract_text() + "\n"
                return text
        
        elif file_extension == '.docx':
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text
        
        else:
            print(f"Unsupported file type: {file_extension}")
            return None
            
    except Exception as e:
        print(f"Error processing document {file_path}: {str(e)}")
        return None


def chunk_text(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[str]:
    """Split text into overlapping chunks for better retrieval."""
    if len(text) <= chunk_size:
        return [text]
    
    chunks = []
    start = 0
    
    while start < len(text):
        # Find the end of the chunk
        end = start + chunk_size
        
        # If we're not at the end of the text, try to break at a sentence or word boundary
        if end < len(text):
            # Look for sentence endings
            sentence_end = text.rfind('.', start, end)
            if sentence_end != -1 and sentence_end > start + chunk_size // 2:
                end = sentence_end + 1
            else:
                # Look for word boundaries
                word_end = text.rfind(' ', start, end)
                if word_end != -1 and word_end > start + chunk_size // 2:
                    end = word_end
        
        chunk = text[start:end].strip()
        if chunk:
            chunks.append(chunk)
        
        # Move start position with overlap
        start = end - chunk_overlap
        if start <= 0:
            start = end
    
    return chunks


def validate_finance_question(question: str) -> bool:
    """Basic validation to check if question is finance-related."""
    finance_keywords = [
        'money', 'finance', 'financial', 'investment', 'invest', 'budget', 'budgeting',
        'savings', 'save', 'debt', 'credit', 'loan', 'mortgage', 'retirement',
        'pension', 'tax', 'taxes', 'income', 'expense', 'spending', 'cost',
        'price', 'stock', 'stocks', 'bond', 'bonds', 'portfolio', 'asset',
        'liability', 'equity', 'fund', 'funds', 'account', 'bank', 'banking',
        'insurance', 'risk', 'return', 'profit', 'loss', 'dividend', 'interest',
        'rate', 'apr', 'apy', 'compound', 'simple', 'principal', 'balance',
        'payment', 'pay', 'owe', 'owing', 'afford', 'emergency', 'goal',
        'wealth', 'rich', 'poor', '401k', 'ira', 'roth', 'traditional',
        # Career and income related
        'salary', 'wage', 'wages', 'negotiate', 'negotiation', 'raise', 'promotion',
        'compensation', 'paycheck', 'bonus', 'benefits', 'career', 'job', 'work',
        'employment', 'employer', 'employee', 'freelance', 'contractor', 'hourly',
        # Real estate and major purchases
        'house', 'home', 'rent', 'renting', 'buying', 'selling', 'property',
        'realtor', 'downpayment', 'closing', 'refinance', 'equity',
        # Additional finance terms
        'networth', 'cashflow', 'frugal', 'cheap', 'expensive', 'value',
        'worth', 'cost', 'price', 'deal', 'bargain', 'discount', 'sale'
    ]
    
    question_lower = question.lower()
    return any(keyword in question_lower for keyword in finance_keywords)


def format_response(response: str) -> str:
    """Format the AI response for better display."""
    # Remove excessive whitespace
    response = ' '.join(response.split())
    
    # Ensure proper sentence endings
    if response and not response.endswith(('.', '!', '?')):
        response += '.'
    
    return response


def get_file_size_mb(file_content: bytes) -> float:
    """Get file size in MB."""
    return len(file_content) / (1024 * 1024) 